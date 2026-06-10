import sys
import os
import json
import docx
import datetime

def fill_docx(data_path, output_path):
    # Load data
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    target_name = data.get("targetName", "Phòng QLDA")
    receiver_name = data.get("receiverName", "Ngô Xuân Việt")
    items = data.get("items", [])
    
    # Path to template
    template_path = os.path.join(os.path.dirname(__file__), "..", "public", "templates", "phieu_cap_phat_vpp.docx")
    doc = docx.Document(template_path)
    
    # Get current date info
    now = datetime.datetime.now()
    day = str(now.day).zfill(2)
    month = str(now.month)
    year = str(now.year)
    
    # 2. Modify paragraphs
    for p in doc.paragraphs:
        # Search and replace metadata
        if "Yêu cầu VPP tháng:" in p.text:
            p.text = p.text.replace("Tháng 6", f"Tháng {month}")
        if "Người yêu cầu:" in p.text:
            p.text = p.text.replace("Ngô Xuân Việt", receiver_name)
        if "Bộ phận:" in p.text:
            p.text = p.text.replace("Phòng QLDA", target_name)
        if "Đề xuất mua các loại văn phòng phẩm" in p.text:
            p.text = p.text.replace("Phòng QLDA", target_name)
        if "TPHCM, ngày 10 tháng 06 năm 2026" in p.text:
            p.text = p.text.replace("ngày 10 tháng 06 năm 2026", f"ngày {day} tháng {month.zfill(2)} năm {year}")

    # 3. Modify table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Fill in items
        for r_idx in range(1, 11):
            if r_idx < len(table.rows) - 1: # avoid header and total row
                row = table.rows[r_idx]
                item_idx = r_idx - 1
                if item_idx < len(items):
                    item = items[item_idx]
                    row.cells[0].text = str(r_idx)
                    row.cells[1].text = item.get("name", "")
                    row.cells[2].text = item.get("unit", "")
                    row.cells[3].text = str(item.get("qty", ""))
                    row.cells[4].text = ""
                    row.cells[5].text = ""
                    row.cells[6].text = item.get("notes", "Đã duyệt cấp phát")
                else:
                    # Clear values for empty rows
                    for c_idx in range(7):
                        row.cells[c_idx].text = ""

    # Save output
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python fill_docx.py <data_path> <output_path>")
        sys.exit(1)
    fill_docx(sys.argv[1], sys.argv[2])
