import sys
import os
import json
import docx
import datetime
import re

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def normalize_dept(dept_name):
    if not dept_name:
        return ""
    d = str(dept_name).lower().strip()
    if "xuyên tâm" in d or "rxt" in d:
        return "rxt"
    if "vàm lẽo" in d or "vàm lẻo" in d:
        return "vam_leo"
    if "mã đà" in d:
        return "ma_da"
    if "trà vinh" in d or "dmt" in d:
        return "dmt_tra_vinh"
    if "tỉnh lộ 8" in d or "tl8" in d:
        return "tl8"
    if "tây ninh" in d:
        return "tay_ninh"
    if "chống hạn" in d:
        return "chong_han"
    if "cà ná" in d:
        return "ca_na"
    if "quản lý dự án" in d or "qlda" in d:
        return "phong_qlda"
    if "hành chính" in d or "hcns" in d or "nhân sự" in d:
        return "phong_hcns"
    if "kế hoạch" in d or "đấu thầu" in d or "kế toán" in d:
        return "phong_hcns" if "hành chính" in d else "phong_ke_hoach"
    if "an toàn" in d or "atld" in d or "atlđ" in d:
        return "phong_atld"
    if "vật tư" in d or "vt-tb" in d or "vt tb" in d or "vt_tb" in d:
        return "phong_vt_tb"
    if "kỹ thuật" in d or "ktht" in d:
        return "phong_ktht"
    return d

def normalize_role(role_name):
    if not role_name:
        return ""
    r = str(role_name).lower().strip()
    if "tuyển dụng" in r:
        return "tuyen_dung"
    if "media" in r:
        return "media"
    if "designer" in r:
        return "designer"
    if "tổ trưởng" in r and "nhân sự" in r:
        return "to_truong_ns"
    if "phó phòng" in r and ("hành chính" in r or "hc" in r):
        return "pp_hanh_chinh"
    if "hành chính" in r or "hành chánh" in r or "hcns" in r or "nhân sự" in r:
        return "hanh_chinh"
    if "an toàn" in r or "atld" in r or "atlđ" in r or "hse" in r or "môi trường" in r:
        return "atld"
    if "qlda" in r or "quản lý dự án" in r or "tổng hợp p qlda" in r:
        return "cv_qlda"
    if "qs" in r:
        return "qs"
    if "đấu thầu" in r:
        return "dau_thau"
    if "kế hoạch" in r or "kinh tế" in r:
        return "ke_hoach"
    if "vật tư" in r or "vt-tb" in r:
        return "vt_tb"
    if "trưởng phòng" in r and "ktht" in r:
        return "tp_ktht"
    if "phó phòng" in r and "ktht" in r:
        return "pp_ktht"
    if "kỹ thuật" in r or "ktht" in r or "kỹ sư" in r:
        return "ktht"
    return r

def parse_date(d_str):
    if not d_str:
        return None
    match = re.search(r"(\d{4})[-/](\d{2})[-/](\d{2})", d_str)
    if match:
        return datetime.date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    return None

def fill_weekly_report(data_path, output_path):
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    start_date_str = data.get("startDate")
    end_date_str = data.get("endDate")
    candidates = data.get("candidates", [])
    office_needs = data.get("officeManualNeeds") or {}
    project_needs = data.get("projectManualNeeds") or {}
    
    start_date = parse_date(start_date_str)
    end_date = parse_date(end_date_str)
    
    print(f"Filtering candidates from {start_date} to {end_date}")
    
    # Filter hired candidates within date range
    hired_candidates = []
    for c in candidates:
        status = str(c.get("status") or "").lower()
        v2_result = str(c.get("v2_result") or "").lower()
        prob_res = str(c.get("probation_result") or "").lower()
        onboard_date_str = c.get("onboard_date")
        
        is_hired = (status in ["hired", "offer"] or 
                    v2_result in ["đạt", "pass"] or 
                    prob_res in ["x", "đạt"] or 
                    onboard_date_str is not None)
        
        if not is_hired:
            continue
            
        c_date_str = onboard_date_str or c.get("v2_interview_date") or c.get("created_at") or c.get("v1_date")
        c_date = parse_date(c_date_str)
        
        if c_date:
            if start_date and c_date < start_date:
                continue
            if end_date and c_date > end_date:
                continue
        
        hired_candidates.append(c)
        
    print(f"Found {len(hired_candidates)} hired candidates in this date range.")

    # Load template
    template_path = os.path.join(os.path.dirname(__file__), "..", "public", "templates", "bao_cao_tuyen_dung_tuan.docx")
    doc = docx.Document(template_path)
    
    # Modify Title Paragraph with date info
    formatted_start = ""
    formatted_end = ""
    if start_date:
        formatted_start = start_date.strftime("%d/%m")
    if end_date:
        formatted_end = end_date.strftime("%d/%m/%Y")
        
    date_range_str = ""
    if formatted_start and formatted_end:
        date_range_str = f" (Từ ngày {formatted_start} đến ngày {formatted_end})"
    
    for p in doc.paragraphs:
        if "BÁO CÁO TUYỂN DỤNG" in p.text or "BÁO CÁO TUYỂN DỤNG THÁNG" in p.text:
            # Modify only the run that contains the title text to preserve other runs (e.g. logo drawing)
            title_updated = False
            for run in p.runs:
                if "BÁO CÁO TUYỂN DỤNG" in run.text or "BÁO CÁO TUYỂN DỤNG THÁNG" in run.text:
                    run.text = f"BÁO CÁO TUYỂN DỤNG TUẦN\n(Từ ngày {formatted_start} đến ngày {formatted_end})"
                    title_updated = True
                    break
            if not title_updated:
                p.text = f"BÁO CÁO TUYỂN DỤNG TUẦN\n(Từ ngày {formatted_start} đến ngày {formatted_end})"

    # Modify bottom date
    report_date = end_date or datetime.date.today()
    day_str = str(report_date.day).zfill(2)
    month_str = str(report_date.month).zfill(2)
    year_str = str(report_date.year)
    
    for p in doc.paragraphs:
        if "ngày" in p.text.lower() and "tháng" in p.text.lower() and "năm" in p.text.lower() and len(p.text) < 50:
            p.text = f"Ngày {day_str} tháng {month_str} năm {year_str}"
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.italic = True

    # Edit Table 0
    if len(doc.tables) > 0:
        table = doc.tables[0]
        
        # Row layout in template
        # Row 0: Header 'Dự án / Phòng ban', 'Vị trí tuyển dụng'...
        # Row 1: 'KHỐI DỰ ÁN'
        # Row 2: Vàm Lẽo header. Row 3, 4: Data rows
        # Row 5: Rạch Xuyên Tâm header. Row 6: Data row
        # Row 7: Tỉnh lộ 8 header. Row 8: Data row
        # Row 9: XLNT Tây Ninh header. Row 10, 11: Data rows
        # Row 12: KCN Cà Ná header. Row 13, 14: Data rows
        # Row 15: KHỐI VĂN PHÒNG header
        # Row 16: P. Kế hoạch đấu thầu header. Row 17, 18: Data rows
        # Row 19: P. Quản lý dự án header. Row 20, 21: Data rows
        # Row 22: P. Hành chính Nhân sự header. Row 23, 24: Data rows
        # Row 25: P. An toàn lao động header. Row 26: Data row
        # Row 27: P. Vật tư thiết bị header. Row 28: Data row
        # Row 29: P. Kỹ thuật header. Row 30, 31: Data rows

        project_slots = [
            {"header_idx": 2, "data_indices": [3, 4]},
            {"header_idx": 5, "data_indices": [6]},
            {"header_idx": 7, "data_indices": [8]},
            {"header_idx": 9, "data_indices": [10, 11]},
            {"header_idx": 12, "data_indices": [13, 14]},
        ]

        office_slots = [
            {"header_idx": 16, "data_indices": [17, 18]},
            {"header_idx": 19, "data_indices": [20, 21]},
            {"header_idx": 22, "data_indices": [23, 24]},
            {"header_idx": 25, "data_indices": [26]},
            {"header_idx": 27, "data_indices": [28]},
            {"header_idx": 29, "data_indices": [30, 31]},
        ]
        
        active_projects = list(project_needs.items())
        active_office = list(office_needs.items())
        
        rows_to_delete = []
        
        # Helper to generate rows for a dept/project
        def generate_rows_for_dept(dept_name, total_demand):
            # Find all candidates in database for this dept (regardless of status/date)
            dept_cands = [c for c in candidates if normalize_dept(c.get("department")) == normalize_dept(dept_name)]
            # Filter hired in range
            hired_in_range = [c for c in hired_candidates if normalize_dept(c.get("department")) == normalize_dept(dept_name)]
            
            rows = []
            if len(hired_in_range) > 0:
                role_groups = {}
                for c in hired_in_range:
                    r_name = c.get("role") or "Nhân sự"
                    if r_name not in role_groups:
                        role_groups[r_name] = []
                    role_groups[r_name].append(c)
                
                hired_total = 0
                for role_name, cands in role_groups.items():
                    hired = len(cands)
                    hired_total += hired
                    # allocate 1 per hired as baseline demand for this role
                    demand = hired
                    
                    notes_parts = []
                    for c in cands:
                        c_name = c.get("name") or ""
                        o_date_str = c.get("onboard_date")
                        if o_date_str:
                            o_date = parse_date(o_date_str)
                            if o_date:
                                notes_parts.append(f"{c_name} {o_date.day}/{o_date.month}")
                            else:
                                notes_parts.append(f"{c_name} {o_date_str}")
                        else:
                            notes_parts.append(c_name)
                            
                    rows.append({
                        "role": role_name,
                        "demand": demand,
                        "hired": hired,
                        "remaining": 0,
                        "progress": "Hoàn thành",
                        "notes": ", ".join(notes_parts)
                    })
                
                if hired_total < total_demand:
                    remaining_demand = total_demand - hired_total
                    other_roles = list(set([c.get("role") for c in dept_cands if c.get("role")]))
                    active_hired_roles = list(role_groups.keys())
                    unused_roles = [r for r in other_roles if r not in active_hired_roles]
                    fallback_role = unused_roles[0] if unused_roles else (other_roles[0] if other_roles else "")
                    
                    rows.append({
                        "role": fallback_role,
                        "demand": remaining_demand,
                        "hired": 0,
                        "remaining": remaining_demand,
                        "progress": "Chưa hoàn thành",
                        "notes": ""
                    })
            else:
                dept_roles = list(set([c.get("role") for c in dept_cands if c.get("role")]))
                primary_role = dept_roles[0] if dept_roles else ""
                
                rows.append({
                    "role": primary_role,
                    "demand": total_demand,
                    "hired": 0,
                    "remaining": total_demand,
                    "progress": "Chưa hoàn thành" if total_demand > 0 else "",
                    "notes": ""
                })
            return rows

        # Process project slots
        for i, slot in enumerate(project_slots):
            if i < len(active_projects):
                name, demand = active_projects[i]
                header_row = table.rows[slot["header_idx"]]
                header_row.cells[0].text = name.upper()
                header_row.cells[0].paragraphs[0].runs[0].font.bold = True
                
                # Generate rows
                dept_rows = generate_rows_for_dept(name, demand)
                
                for j, r_data in enumerate(dept_rows):
                    if j < len(slot["data_indices"]):
                        data_row = table.rows[slot["data_indices"][j]]
                        
                        data_row.cells[0].text = "" # Dept name (merged/left empty)
                        data_row.cells[1].text = r_data["role"]
                        data_row.cells[2].text = str(r_data["demand"])
                        data_row.cells[3].text = str(r_data["hired"]) if r_data["hired"] > 0 else "-"
                        
                        if r_data["demand"] > 0:
                            data_row.cells[4].text = str(r_data["remaining"]) if r_data["remaining"] > 0 else "0"
                            data_row.cells[5].text = r_data["progress"]
                        else:
                            data_row.cells[4].text = "-"
                            data_row.cells[5].text = ""
                        data_row.cells[6].text = r_data["notes"]
                    else:
                        # If more rows generated than slot template rows, we append them or ignore
                        # For this app, templates have enough rows
                        pass
                
                # Mark unused data rows in this slot for deletion
                if len(dept_rows) < len(slot["data_indices"]):
                    rows_to_delete.extend(slot["data_indices"][len(dept_rows):])
            else:
                # Unused slot - delete header and all data rows
                rows_to_delete.append(slot["header_idx"])
                rows_to_delete.extend(slot["data_indices"])

        # Process office slots
        for i, slot in enumerate(office_slots):
            if i < len(active_office):
                name, demand = active_office[i]
                header_row = table.rows[slot["header_idx"]]
                
                display_name = name
                if name == "HCNS":
                    display_name = "PHÒNG HÀNH CHÍNH - NHÂN SỰ"
                elif name == "Phòng QLDA":
                    display_name = "PHÒNG QUẢN LÝ DỰ ÁN"
                elif name == "Kế toán":
                    display_name = "PHÒNG KẾ TOÁN"
                
                header_row.cells[0].text = display_name.upper()
                header_row.cells[0].paragraphs[0].runs[0].font.bold = True
                
                # Generate rows
                dept_rows = generate_rows_for_dept(name, demand)
                
                for j, r_data in enumerate(dept_rows):
                    if j < len(slot["data_indices"]):
                        data_row = table.rows[slot["data_indices"][j]]
                        
                        data_row.cells[0].text = ""
                        data_row.cells[1].text = r_data["role"]
                        data_row.cells[2].text = str(r_data["demand"])
                        data_row.cells[3].text = str(r_data["hired"]) if r_data["hired"] > 0 else "-"
                        
                        if r_data["demand"] > 0:
                            data_row.cells[4].text = str(r_data["remaining"]) if r_data["remaining"] > 0 else "0"
                            data_row.cells[5].text = r_data["progress"]
                        else:
                            data_row.cells[4].text = "-"
                            data_row.cells[5].text = ""
                        data_row.cells[6].text = r_data["notes"]
                
                if len(dept_rows) < len(slot["data_indices"]):
                    rows_to_delete.extend(slot["data_indices"][len(dept_rows):])
            else:
                rows_to_delete.append(slot["header_idx"])
                rows_to_delete.extend(slot["data_indices"])

        # Delete rows in descending order
        rows_to_delete = sorted(list(set(rows_to_delete)), reverse=True)
        for r_idx in rows_to_delete:
            tr = table.rows[r_idx]._tr
            table._tbl.remove(tr)

    doc.save(output_path)
    print(f"Report successfully saved to {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python fill_weekly_report.py <data_path> <output_path>")
        sys.exit(1)
    fill_weekly_report(sys.argv[1], sys.argv[2])
